from __future__ import annotations

import io
import tempfile
import threading
import time
from datetime import datetime
from pathlib import Path

import pytest
from fastapi.testclient import TestClient
import httpx
from qdrant_client import QdrantClient, models

try:  # pragma: no cover - optional dependency
    from pypdf import PdfReader  # noqa: F401
    HAS_PYPDF = True
except Exception:  # pragma: no cover
    PdfReader = None  # type: ignore
    HAS_PYPDF = False

try:  # pragma: no cover - optional dependency
    from docx import Document as DocxDocument
    HAS_DOCX = True
except Exception:  # pragma: no cover
    DocxDocument = None  # type: ignore
    HAS_DOCX = False

try:  # pragma: no cover - optional dependency
    from bs4 import BeautifulSoup  # noqa: F401
    HAS_BS4 = True
except Exception:  # pragma: no cover
    HAS_BS4 = False
HAS_PDF = HAS_PYPDF

from cornerstone.app import create_app
from cornerstone.chat import SupportAgentService
from cornerstone.config import Settings
from cornerstone.glossary import Glossary
from cornerstone.ingestion import (
    DocumentIngestor,
    IngestionJobManager,
    IngestionResult,
    ProjectVectorStoreManager,
)
from cornerstone.fts import FTSIndex
from cornerstone.personas import PersonaStore
from cornerstone.projects import DocumentMetadata, ProjectStore


class FakeEmbeddingService:
    def __init__(self) -> None:
        self.dimension = 3

    def embed(self, texts):
        return [[float(len(text)), 0.0, 0.0] for text in texts]

    def embed_one(self, text: str):
        return [float(len(text)), 0.0, 0.0]


# Minimal PDF with extractable text ("PDF sample text")
SAMPLE_PDF_BYTES = b"%PDF-1.1\n1 0 obj<<>>endobj\n2 0 obj<< /Length 56 >>stream\nBT /F1 12 Tf 72 720 Td (PDF sample text) Tj ET\nendstream\nendobj\n3 0 obj<< /Type /Page /Parent 4 0 R /MediaBox [0 0 612 792] /Contents 2 0 R /Resources << /Font << /F1 5 0 R >> >> >>endobj\n4 0 obj<< /Type /Pages /Kids [3 0 R] /Count 1 >>endobj\n5 0 obj<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>endobj\n6 0 obj<< /Type /Catalog /Pages 4 0 R >>endobj\nxref\n0 7\n0000000000 65535 f \n0000000010 00000 n \n0000000056 00000 n \n0000000125 00000 n \n0000000230 00000 n \n0000000302 00000 n \n0000000373 00000 n \ntrailer<< /Root 6 0 R /Size 7 >>\nstartxref\n430\n%%EOF"

HTML_SAMPLE_BYTES = b"""<html><head><title>Support FAQ</title></head><body><h1>FAQ</h1><p>Step one.</p></body></html>"""


def build_app(settings: Settings | None = None):
    tmpdir = Path(tempfile.mkdtemp(prefix="cornerstone-ingest-"))
    base_settings = settings or Settings()
    settings_kwargs = {field: getattr(base_settings, field) for field in Settings.__dataclass_fields__ if field != 'local_data_dir'}
    settings_kwargs.update(
        {
            "data_dir": str(tmpdir),
            "local_data_dir": str(Path(tmpdir) / "local"),
        }
    )
    settings = Settings(**settings_kwargs)
    Path(settings.local_data_dir).mkdir(parents=True, exist_ok=True)
    embedding = FakeEmbeddingService()
    project_store = ProjectStore(tmpdir, default_project_name=settings.default_project_name)
    persona_store = PersonaStore(tmpdir)
    default_project = project_store.list_projects()[0]

    client = QdrantClient(path=":memory:")
    store_manager = ProjectVectorStoreManager(
        client_factory=lambda: client,
        vector_size=embedding.dimension,
        distance=models.Distance.COSINE,
        collection_name_fn=settings.project_collection_name,
    )
    store_manager.get_store(default_project.id)

    glossary = Glossary()
    fts_index = FTSIndex(Path(tmpdir) / "fts.sqlite")
    chat_service = SupportAgentService(
        settings=settings,
        embedding_service=embedding,  # type: ignore[arg-type]
        store_manager=store_manager,
        glossary=glossary,
        persona_store=persona_store,
        fts_index=fts_index,
    )
    ingestion_service = DocumentIngestor(embedding, store_manager, project_store, fts_index=fts_index)

    app = create_app(
        settings=settings,
        embedding_service=embedding,  # type: ignore[arg-type]
        glossary=glossary,
        project_store=project_store,
        persona_store=persona_store,
        store_manager=store_manager,
        chat_service=chat_service,
        ingestion_service=ingestion_service,
    )
    app.state.services.fts_index = fts_index

    return TestClient(app)


def _build_docx_bytes() -> bytes:
    if not HAS_DOCX:
        pytest.importorskip("docx")
    document = DocxDocument()
    document.add_heading("Docx Title", level=1)
    document.add_paragraph("Body text content for testing")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_uploads_across_multiple_projects():
    client = build_app()
    state = client.app.state.services
    project_store: ProjectStore = state.project_store

    # Create two additional projects via API
    resp = client.post(
        "/knowledge/projects",
        data={"name": "Project Two", "description": ""},
        follow_redirects=False,
    )
    assert resp.status_code == 303
    resp = client.post(
        "/knowledge/projects",
        data={"name": "Project Three", "description": "Extra"},
        follow_redirects=False,
    )
    assert resp.status_code == 303

    projects = project_store.list_projects()
    assert len(projects) >= 3

    uploads = [
        (projects[0], ("notes.txt", b"Alpha troubleshooting steps", "text/plain")),
        (projects[1], ("guide.md", b"# Beta issue\n\n## Fix", "text/markdown")),
        (
            projects[2],
            (
                "manual.pdf" if HAS_PYPDF else "manual.txt",
                SAMPLE_PDF_BYTES if HAS_PYPDF else b"PDF fallback text",
                "application/pdf" if HAS_PYPDF else "text/plain",
            ),
        ),
    ]

    for project, file_tuple in uploads:
        response = client.post(
            "/knowledge/upload",
            data={"project_id": project.id},
            files={"file": file_tuple},
            follow_redirects=False,
        )
        assert response.status_code == 303, response.text

    for project, file_tuple in uploads:
        docs = project_store.list_documents(project.id)
        assert len(docs) == 1
        metadata = docs[0]
        assert metadata.filename == file_tuple[0]
        assert metadata.chunk_count > 0
        assert metadata.content_type == file_tuple[2]
        if file_tuple[0].endswith('.md'):
            assert metadata.title == 'Beta issue'
        elif file_tuple[0] == 'notes.txt':
            assert metadata.title == 'Alpha troubleshooting steps'
        elif file_tuple[0] == 'manual.txt':
            assert metadata.title == 'PDF fallback text'
        elif file_tuple[0].endswith('.pdf') and HAS_PYPDF:
            assert metadata.title

        store = state.store_manager.get_store(project.id)
        assert store.count() >= metadata.chunk_count
        results = store.search(state.embedding_service.embed_one("sanity"), limit=1)
        assert results
        assert all(result.payload.get("project_id") == project.id for result in results if result.payload)

    # Delete the second project's document and ensure cleanup
    target_project, _ = uploads[1]
    target_doc = project_store.list_documents(target_project.id)[0]
    response = client.post(
        "/knowledge/delete",
        data={"project_id": target_project.id, "doc_id": target_doc.id},
        follow_redirects=False,
    )
    assert response.status_code == 303
    assert project_store.list_documents(target_project.id) == []
    store = state.store_manager.get_store(target_project.id)
    assert store.count() == 0


def test_fts_index_populated_after_ingest():
    client = build_app()
    state = client.app.state.services
    project = state.project_store.list_projects()[0]

    response = client.post(
        "/knowledge/upload",
        data={"project_id": project.id},
        files={"file": ("alpha.md", b"# Alpha\nContent about Alpha", "text/markdown")},
        follow_redirects=False,
    )
    assert response.status_code == 303

    hits = state.fts_index.search(project.id, "Alpha", limit=2)
    assert hits
    assert any("Alpha" in hit["text"] for hit in hits)


def test_async_multi_file_uploads_return_jobs():
    client = build_app()
    state = client.app.state.services
    project = state.project_store.list_projects()[0]

    files = [
        (
            'files',
            ('alpha.txt', b'Alpha body', 'text/plain'),
        ),
        (
            'files',
            ('bravo.md', b'# Bravo\nMore details', 'text/markdown'),
        ),
    ]

    response = client.post(
        "/knowledge/uploads",
        data={"project_id": project.id},
        files=files,
    )
    assert response.status_code == 202
    payload = response.json()
    assert len(payload["jobs"]) == 2
    job_ids = {job["id"] for job in payload["jobs"]}
    assert job_ids

    for _ in range(40):
        poll = client.get("/knowledge/uploads", params={"project_id": project.id})
        poll.raise_for_status()
        jobs = poll.json().get("jobs", [])
        tracked = [job for job in jobs if job["id"] in job_ids]
        if tracked and all(job["status"] in {"completed", "failed"} for job in tracked):
            break
        time.sleep(0.05)
    else:
        pytest.fail("Ingestion jobs did not finish in time")

    docs = state.project_store.list_documents(project.id)
    filenames = {doc.filename for doc in docs}
    assert {"alpha.txt", "bravo.md"}.issubset(filenames)


def test_metadata_enrichment_for_docx_and_html():
    client = build_app()
    state = client.app.state.services
    project = state.project_store.list_projects()[0]

    if not HAS_BS4:
        pytest.skip("beautifulsoup4 is required for HTML extraction")

    docx_bytes = _build_docx_bytes()
    html_bytes = HTML_SAMPLE_BYTES

    files = [
        (
            'files',
            (
                'analysis.docx',
                docx_bytes,
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            ),
        ),
        (
            'files',
            ('page.html', html_bytes, 'text/html'),
        ),
    ]

    response = client.post(
        "/knowledge/uploads",
        data={"project_id": project.id},
        files=files,
    )
    assert response.status_code == 202

    for _ in range(40):
        poll = client.get("/knowledge/uploads", params={"project_id": project.id})
        poll.raise_for_status()
        jobs = poll.json().get("jobs", [])
        if jobs and all(job["status"] == "completed" for job in jobs):
            break
        time.sleep(0.05)
    else:
        pytest.fail("Docx/HTML ingestion jobs did not finish")

    documents = sorted(state.project_store.list_documents(project.id), key=lambda meta: meta.filename)
    assert documents[0].filename == 'analysis.docx'
    assert documents[0].title == 'Docx Title'
    assert documents[0].content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    assert documents[1].filename == 'page.html'
    if HAS_BS4:
        assert documents[1].title == 'Support FAQ'
        assert documents[1].content_type == 'text/html'


def test_ingest_url_downloads_and_extracts(monkeypatch):
    if not HAS_BS4:
        pytest.importorskip("bs4")

    client = build_app()
    state = client.app.state.services
    ingestion: DocumentIngestor = state.ingestion_service
    project = state.project_store.list_projects()[0]

    class DummyResponse:
        status_code = 200
        headers = {"content-type": "text/html; charset=utf-8"}
        content = HTML_SAMPLE_BYTES

    monkeypatch.setattr(httpx, "get", lambda url, timeout=10.0, follow_redirects=True: DummyResponse())

    result = ingestion.ingest_url(project.id, url="https://example.com/docs")
    assert result.document.filename.endswith('.html')
    assert result.document.content_type == 'text/html'
    assert result.document.title == 'Support FAQ'


def test_upload_url_endpoint_enqueues_job(monkeypatch):
    client = build_app()
    state = client.app.state.services
    project = state.project_store.list_projects()[0]

    fake_document = DocumentMetadata(
        id="doc123",
        filename="remote.html",
        chunk_count=2,
        created_at=datetime.now().isoformat(),
        size_bytes=42,
        title="Remote Doc",
        content_type="text/html",
    )

    def fake_ingest_url(project_id: str, *, url: str, timeout: float = 10.0) -> IngestionResult:
        assert project_id == project.id
        assert url == "https://example.com/help"
        return IngestionResult(document=fake_document, chunks_ingested=2)

    monkeypatch.setattr(state.ingestion_service, "ingest_url", fake_ingest_url)

    response = client.post(
        "/knowledge/upload-url",
        data={"project_id": project.id, "url": "https://example.com/help"},
    )
    assert response.status_code == 202
    job = response.json()["job"]
    assert job["status"] == "pending"

    poll = client.get("/knowledge/uploads", params={"project_id": project.id})
    poll.raise_for_status()
    jobs = poll.json().get("jobs", [])
    assert any(j["status"] == "completed" and j.get("document", {}).get("filename") == "remote.html" for j in jobs)


@pytest.mark.parametrize(
    "filename, factory, content_type, expected_title",
    [
        (
            "plain.txt",
            lambda: b"Plain heading\nMore text",
            "text/plain",
            "Plain heading",
        ),
        (
            "guide.md",
            lambda: b"# Markdown Heading\n\nDetails",
            "text/markdown",
            "Markdown Heading",
        ),
        (
            "manual.pdf",
            lambda: SAMPLE_PDF_BYTES if HAS_PDF else b"PDF fallback text",
            "application/pdf" if HAS_PDF else "text/plain",
            "PDF sample text" if HAS_PDF else "PDF fallback text",
        ),
        (
            "analysis.docx",
            _build_docx_bytes,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "Docx Title",
        ),
        (
            "faq.html",
            lambda: HTML_SAMPLE_BYTES,
            "text/html",
            "Support FAQ",
        ),
    ],
)
def test_single_upload_supported_types(filename, factory, content_type, expected_title):
    if filename.endswith(".docx") and not HAS_DOCX:
        pytest.skip("python-docx is required for DOCX ingestion")
    if filename.endswith(".html") and not HAS_BS4:
        pytest.skip("beautifulsoup4 is required for HTML ingestion")
    if filename.endswith(".pdf") and not HAS_PDF:
        pytest.skip("pypdf is required for PDF ingestion")

    client = build_app()
    state = client.app.state.services
    project = state.project_store.list_projects()[0]

    payload = {
        "project_id": project.id,
    }
    file_bytes = factory() if callable(factory) else factory

    response = client.post(
        "/knowledge/upload",
        data=payload,
        files={"file": (filename, file_bytes, content_type)},
        follow_redirects=False,
    )
    assert response.status_code == 303

    documents = state.project_store.list_documents(project.id)
    assert len(documents) == 1
    doc = documents[0]
    assert doc.filename == filename
    assert doc.chunk_count > 0
    assert doc.content_type == content_type
    if expected_title:
        assert doc.title == expected_title
    else:
        assert doc.title is None


def test_cleanup_endpoint_purges_project_vectors():
    client = build_app()
    state = client.app.state.services
    project_store: ProjectStore = state.project_store
    project = project_store.list_projects()[0]

    response = client.post(
        "/knowledge/upload",
        data={"project_id": project.id},
        files={"file": ("guide.txt", b"Troubleshooting steps", "text/plain")},
        follow_redirects=False,
    )
    assert response.status_code == 303

    documents = project_store.list_documents(project.id)
    assert documents
    store = state.store_manager.get_store(project.id)
    assert store.count() > 0

    response = client.post(
        "/knowledge/cleanup",
        data={"project_id": project.id},
        follow_redirects=False,
    )
    assert response.status_code == 303
    assert project_store.list_documents(project.id) == []
    assert store.count() == 0


def test_ingested_chunks_include_metadata_summary_and_language():
    client = build_app()
    state = client.app.state.services
    project = state.project_store.list_projects()[0]

    rich_text = "# Getting Started\n\n" + " ".join(
        f"This sentence number {i} explains the setup." for i in range(120)
    )

    response = client.post(
        "/knowledge/upload",
        data={"project_id": project.id},
        files={"file": ("guide.md", rich_text.encode("utf-8"), "text/markdown")},
        follow_redirects=False,
    )
    assert response.status_code == 303

    payloads = list(state.store_manager.iter_project_payloads(project.id))
    assert payloads
    first_payload = payloads[0]
    assert first_payload.get("summary")
    assert first_payload.get("language") == "en"
    assert first_payload.get("token_count", 0) > 0
    assert "section_path" in first_payload


def test_local_directories_response_includes_stats(tmp_path: Path):
    settings = Settings(
        data_dir=str(tmp_path),
        default_project_name="Project One",
        local_data_dir=str(tmp_path / "local"),
    )
    client = build_app(settings)

    local_root = Path(client.app.state.services.settings.local_data_dir)
    (local_root / "alpha").mkdir(parents=True)
    sample_file = local_root / "alpha" / "guide.md"
    sample_file.write_text("Hello world", encoding="utf-8")
    (local_root / "alpha" / "ignore.bin").write_bytes(b"\x00\x01")

    response = client.get("/knowledge/local-directories")
    assert response.status_code == 200
    data = response.json()
    assert "directories" in data
    alpha = next((entry for entry in data["directories"] if entry["name"] == "alpha"), None)
    assert alpha is not None
    assert alpha.get("path") == "alpha"
    assert alpha.get("file_count") == 1
    assert isinstance(alpha.get("total_bytes"), int)
    assert alpha["total_bytes"] >= len("Hello world".encode("utf-8"))


def test_job_manager_throttles_when_project_over_limit():
    manager = IngestionJobManager(
        max_active_per_project=1,
        max_files_per_minute=None,
        throttle_poll_interval=0.05,
    )
    job1 = manager.create_job("proj", "alpha.txt")
    job2 = manager.create_job("proj", "bravo.txt")

    release = threading.Event()

    def worker(job_id: str, doc_id: str) -> None:
        manager.mark_processing(job_id)
        release.wait()
        manager.mark_completed(
            job_id,
            DocumentMetadata(
                id=doc_id,
                filename="doc.txt",
                chunk_count=1,
                created_at=datetime.now().isoformat(),
                size_bytes=10,
                title="Doc",
                content_type="text/plain",
            ),
        )

    t1 = threading.Thread(target=worker, args=(job1.id, "d1"), daemon=True)
    t2 = threading.Thread(target=worker, args=(job2.id, "d2"), daemon=True)

    t1.start()
    time.sleep(0.05)
    t2.start()

    time.sleep(0.15)
    throttled_job = manager.get(job2.id)
    assert throttled_job is not None
    assert throttled_job.status == "throttled"

    release.set()
    t1.join(timeout=1.0)
    t2.join(timeout=1.0)

    assert manager.get(job1.id).status == "completed"
    assert manager.get(job2.id).status == "completed"


def test_job_manager_wait_for_rate_limits_throughput():
    manager = IngestionJobManager(
        max_active_per_project=2,
        max_files_per_minute=2,
        throttle_poll_interval=0.01,
        rate_window_seconds=1.0,
    )

    start = time.monotonic()
    manager.wait_for_rate("proj")
    manager.wait_for_rate("proj")
    pre_third = time.monotonic() - start
    manager.wait_for_rate("proj")
    total_elapsed = time.monotonic() - start

    assert pre_third < 1.0
    assert total_elapsed >= 0.95


def test_local_import_reports_progress(tmp_path: Path):
    settings = Settings(data_dir=str(tmp_path), default_project_name="Project One")
    client = build_app(settings)
    project = client.app.state.services.project_store.list_projects()[0]
    local_root = Path(client.app.state.services.settings.local_data_dir)
    target = local_root / "docs"
    target.mkdir(parents=True)
    for idx in range(3):
        (target / f"file-{idx}.md").write_text("content", encoding="utf-8")

    response = client.post(
        "/knowledge/local-import",
        data={"project_id": project.id, "path": "docs"},
        follow_redirects=False,
    )
    assert response.status_code == 202
    job = response.json()["job"]
    job_id = job["id"]

    for _ in range(10):
        poll = client.get("/knowledge/uploads", params={"project_id": project.id})
        assert poll.status_code == 200
        jobs = poll.json().get("jobs", [])
        target_job = next((item for item in jobs if item["id"] == job_id), None)
        assert target_job is not None
        if target_job["status"] == "completed":
            assert target_job.get("total_files") == 3
            assert target_job.get("processed_files") == 3
            assert target_job.get("processed_bytes", 0) >= 0
            break
    else:  # pragma: no cover - defensive timeout
        pytest.fail("Local import job did not complete")
