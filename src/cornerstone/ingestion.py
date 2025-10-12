"""Document ingestion utilities for project-specific knowledge bases."""

from __future__ import annotations

import io
import time
from collections import defaultdict, deque
from dataclasses import dataclass, asdict
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable, Sequence
from urllib.parse import urlparse
from uuid import uuid4

import httpx
from fastapi import UploadFile

try:  # pragma: no cover - optional dependency
    from pypdf import PdfReader
except Exception:  # pragma: no cover
    PdfReader = None

try:  # pragma: no cover - optional dependency
    from docx import Document as DocxDocument
except Exception:  # pragma: no cover
    DocxDocument = None

try:  # pragma: no cover - optional dependency
    from bs4 import BeautifulSoup
except Exception:  # pragma: no cover
    BeautifulSoup = None

import logging
import mimetypes
import threading

from .embeddings import EmbeddingService
from .projects import DocumentMetadata, ProjectStore
from .vector_store import QdrantVectorStore, VectorRecord
from .chunker import chunk_text, Chunk
from .fts import FTSIndex
from .observability import MetricsRecorder
from qdrant_client import models


logger = logging.getLogger(__name__)


def _require_beautifulsoup():
    global BeautifulSoup  # type: ignore[assignment]
    if BeautifulSoup is not None:
        return BeautifulSoup
    try:  # pragma: no cover - optional import at runtime
        from bs4 import BeautifulSoup as _BeautifulSoup

        BeautifulSoup = _BeautifulSoup  # type: ignore[assignment]
        return BeautifulSoup
    except Exception:
        return None


@dataclass(slots=True)
class IngestionResult:
    document: DocumentMetadata
    chunks_ingested: int


@dataclass(slots=True)
class ExtractedDocument:
    text: str
    title: str | None
    content_type: str | None


@dataclass(slots=True)
class IngestionJob:
    id: str
    project_id: str
    filename: str
    status: str
    created_at: str
    updated_at: str
    document: DocumentMetadata | None = None
    error: str | None = None
    total_files: int | None = None
    processed_files: int | None = None
    total_bytes: int | None = None
    processed_bytes: int | None = None

    def to_dict(self) -> dict:
        return {
            "id": self.id,
            "project_id": self.project_id,
            "filename": self.filename,
            "status": self.status,
            "created_at": self.created_at,
            "updated_at": self.updated_at,
            "error": self.error,
            "document": asdict(self.document) if self.document else None,
            "total_files": self.total_files,
            "processed_files": self.processed_files,
            "total_bytes": self.total_bytes,
            "processed_bytes": self.processed_bytes,
        }


class IngestionJobManager:
    """Track ingestion job status for asynchronous uploads."""

    def __init__(
        self,
        *,
        max_active_per_project: int = 3,
        max_files_per_minute: int | None = 180,
        throttle_poll_interval: float = 0.25,
        rate_window_seconds: float = 60.0,
    ) -> None:
        self._jobs: dict[str, IngestionJob] = {}
        self._lock = threading.Lock()
        self._project_active: dict[str, int] = defaultdict(int)
        self._max_active_per_project = max(1, max_active_per_project)
        self._throttle_poll_interval = max(0.05, throttle_poll_interval)
        self._max_files_per_minute = max_files_per_minute
        self._project_file_times: dict[str, deque[float]] = defaultdict(deque)
        self._rate_window_seconds = max(1.0, rate_window_seconds)

    def create_job(self, project_id: str, filename: str) -> IngestionJob:
        now = self._now()
        job = IngestionJob(
            id=uuid4().hex,
            project_id=project_id,
            filename=filename,
            status="pending",
            created_at=now,
            updated_at=now,
        )
        with self._lock:
            self._jobs[job.id] = job
        return job

    def mark_processing(
        self,
        job_id: str,
        *,
        total_files: int | None = None,
        processed_files: int | None = None,
        total_bytes: int | None = None,
        processed_bytes: int | None = None,
    ) -> None:
        while True:
            with self._lock:
                job = self._jobs.get(job_id)
                if not job:
                    return
                active = self._project_active[job.project_id]
                if active < self._max_active_per_project:
                    job.status = "processing"
                    job.updated_at = self._now()
                    self._project_active[job.project_id] = active + 1
                    if total_files is not None:
                        job.total_files = total_files
                    if processed_files is not None:
                        job.processed_files = processed_files
                    if total_bytes is not None:
                        job.total_bytes = total_bytes
                    if processed_bytes is not None:
                        job.processed_bytes = processed_bytes
                    return
                job.status = "throttled"
                job.updated_at = self._now()
                if processed_files is not None:
                    job.processed_files = processed_files
                if total_files is not None:
                    job.total_files = total_files
            time.sleep(self._throttle_poll_interval)

    def mark_completed(
        self,
        job_id: str,
        document: DocumentMetadata,
        *,
        processed_files: int | None = None,
        processed_bytes: int | None = None,
    ) -> None:
        with self._lock:
            job = self._jobs.get(job_id)
            if not job:
                return
            job.status = "completed"
            job.document = document
            job.error = None
            job.updated_at = self._now()
            if processed_files is not None:
                job.processed_files = processed_files
            if processed_bytes is not None:
                job.processed_bytes = processed_bytes
                if job.total_bytes is None:
                    job.total_bytes = processed_bytes
            if processed_files is not None and job.total_files is None:
                job.total_files = processed_files
            if self._project_active[job.project_id] > 0:
                self._project_active[job.project_id] -= 1

    def mark_failed(
        self,
        job_id: str,
        error: str,
        *,
        processed_files: int | None = None,
        total_files: int | None = None,
    ) -> None:
        with self._lock:
            job = self._jobs.get(job_id)
            if not job:
                return
            job.status = "failed"
            job.error = error
            job.updated_at = self._now()
            if processed_files is not None:
                job.processed_files = processed_files
            if total_files is not None:
                job.total_files = total_files
            if self._project_active[job.project_id] > 0:
                self._project_active[job.project_id] -= 1

    def update_progress(
        self,
        job_id: str,
        *,
        processed_files: int | None = None,
        total_files: int | None = None,
        processed_bytes: int | None = None,
        total_bytes: int | None = None,
        status: str | None = None,
    ) -> None:
        with self._lock:
            job = self._jobs.get(job_id)
            if not job:
                return
            if processed_files is not None:
                job.processed_files = processed_files
            if total_files is not None:
                job.total_files = total_files
            if processed_bytes is not None:
                job.processed_bytes = processed_bytes
            if total_bytes is not None:
                job.total_bytes = total_bytes
            if status is not None:
                job.status = status
            job.updated_at = self._now()

    def list_for_project(self, project_id: str) -> list[IngestionJob]:
        with self._lock:
            return sorted(
                (job for job in self._jobs.values() if job.project_id == project_id),
                key=lambda job: job.created_at,
            )

    def get(self, job_id: str) -> IngestionJob | None:
        with self._lock:
            return self._jobs.get(job_id)

    def wait_for_rate(self, project_id: str) -> None:
        if not self._max_files_per_minute:
            return
        window = self._rate_window_seconds
        limit = max(1, self._max_files_per_minute)
        while True:
            with self._lock:
                history = self._project_file_times[project_id]
                now = time.monotonic()
                while history and now - history[0] > window:
                    history.popleft()
                if len(history) < limit:
                    history.append(now)
                    return
                wait_time = window - (now - history[0])
            time.sleep(max(wait_time, self._throttle_poll_interval))

    @staticmethod
    def _now() -> str:
        return datetime.now(timezone.utc).isoformat()


class ProjectVectorStoreManager:
    """Cache and manage vector stores per project."""

    def __init__(
        self,
        client_factory,
        *,
        vector_size: int,
        distance,
        collection_name_fn,
        collection_kwargs: dict[str, Any] | None = None,
    ) -> None:
        self._client_factory = client_factory
        self._vector_size = vector_size
        self._distance = distance
        self._collection_name_fn = collection_name_fn
        self._collection_kwargs = dict(collection_kwargs or {})
        self._stores: dict[str, QdrantVectorStore] = {}

    def get_store(self, project_id: str) -> QdrantVectorStore:
        if project_id in self._stores:
            return self._stores[project_id]
        collection_name = self._collection_name_fn(project_id)
        store = QdrantVectorStore(
            client=self._client_factory(),
            collection_name=collection_name,
            vector_size=self._vector_size,
            distance=self._distance,
            **self._collection_kwargs,
        )
        store.ensure_collection()
        store.ensure_payload_indexes()
        self._stores[project_id] = store
        return store

    def delete_document(self, project_id: str, doc_id: str) -> bool:
        store = self.get_store(project_id)
        flt = models.Filter(must=[models.FieldCondition(key="doc_id", match=models.MatchValue(value=doc_id))])
        result = store.delete_by_filter(flt)
        return result.status == models.UpdateStatus.COMPLETED

    def purge_project(self, project_id: str) -> bool:
        """Remove all vectors associated with a project."""

        store = self.get_store(project_id)
        flt = models.Filter(
            must=[models.FieldCondition(key="project_id", match=models.MatchValue(value=project_id))]
        )
        result = store.delete_by_filter(flt)
        return result.status == models.UpdateStatus.COMPLETED

    def iter_project_payloads(self, project_id: str, *, batch_size: int = 256):
        """Yield payload dictionaries for all vectors stored for a project."""

        store = self.get_store(project_id)
        flt = models.Filter(
            must=[models.FieldCondition(key="project_id", match=models.MatchValue(value=project_id))]
        )
        yield from store.iter_payloads(scroll_filter=flt, batch_size=batch_size)


class DocumentIngestor:
    """Convert uploaded documents into embeddings stored per project."""

    def __init__(
        self,
        embedding_service: EmbeddingService,
        store_manager: ProjectVectorStoreManager,
        project_store: ProjectStore,
        fts_index: FTSIndex | None = None,
        metrics: MetricsRecorder | None = None,
    ) -> None:
        self._embedding = embedding_service
        self._stores = store_manager
        self._projects = project_store
        self._fts = fts_index
        self._metrics = metrics

    @property
    def embedding_model_id(self) -> str | None:
        """Return the identifier for the active embedding model, if available."""

        return getattr(self._embedding, "model_identifier", None)

    async def ingest_upload(self, project_id: str, upload: UploadFile) -> IngestionResult:
        filename = upload.filename or "document"
        logger.info("ingest.start project=%s filename=%s", project_id, filename)
        raw_bytes = await upload.read()
        return self.ingest_bytes(
            project_id,
            filename=filename,
            data=raw_bytes,
            content_type=upload.content_type,
        )

    def ingest_bytes(
        self,
        project_id: str,
        *,
        filename: str,
        data: bytes,
        content_type: str | None = None,
    ) -> IngestionResult:
        if not data:
            raise ValueError("Uploaded file is empty")

        metrics = self._metrics
        overall_start = time.perf_counter() if metrics else None
        extracted = self._extract_document(filename, data, content_type)
        chunk_objects = chunk_text(
            extracted.text,
            content_type=extracted.content_type,
        )
        if not chunk_objects:
            raise ValueError("No textual content could be extracted from the document")

        doc_id = uuid4().hex
        embed_start = time.perf_counter() if metrics else None
        embeddings = self._embedding.embed([chunk.text for chunk in chunk_objects])
        if metrics and embed_start is not None:
            metrics.record_timing(
                "ingestion.embedding_duration",
                time.perf_counter() - embed_start,
                project_id=project_id,
                chunks=len(chunk_objects),
            )
        store = self._stores.get_store(project_id)
        records, fts_entries = self._build_records(
            doc_id,
            filename,
            chunk_objects,
            embeddings,
            project_id,
            extracted.content_type,
            extracted.title,
        )
        upsert_start = time.perf_counter() if metrics else None
        store.upsert(records)
        if metrics and upsert_start is not None:
            metrics.record_timing(
                "ingestion.vector_upsert_duration",
                time.perf_counter() - upsert_start,
                project_id=project_id,
                chunks=len(records),
            )
        if self._fts is not None:
            fts_start = time.perf_counter() if metrics else None
            self._fts.upsert_chunks(
                project_id=project_id,
                doc_id=doc_id,
                entries=fts_entries,
            )
            if metrics and fts_start is not None:
                metrics.record_timing(
                    "ingestion.fts_upsert_duration",
                    time.perf_counter() - fts_start,
                    project_id=project_id,
                    chunks=len(fts_entries),
                )
        logger.info(
            "ingest.upserted project=%s doc_id=%s chunks=%s",
            project_id,
            doc_id,
            len(records),
        )

        metadata = DocumentMetadata(
            id=doc_id,
            filename=filename,
            chunk_count=len(records),
            created_at=self._now(),
            size_bytes=len(data),
            title=extracted.title,
            content_type=extracted.content_type,
        )
        self._projects.record_document(project_id, metadata)
        logger.info("ingest.completed project=%s doc_id=%s", project_id, doc_id)
        if metrics and overall_start is not None:
            metrics.record_timing(
                "ingestion.total_duration",
                time.perf_counter() - overall_start,
                project_id=project_id,
                content_type=extracted.content_type or "unknown",
                chunks=len(records),
                size_bytes=len(data),
            )
            metrics.increment(
                "ingestion.documents",
                project_id=project_id,
            )
            metrics.increment(
                "ingestion.chunks",
                value=len(records),
                project_id=project_id,
            )
        return IngestionResult(document=metadata, chunks_ingested=len(records))

    def ingest_url(
        self,
        project_id: str,
        *,
        url: str,
        timeout: float = 10.0,
    ) -> IngestionResult:
        parsed = urlparse(url)
        if parsed.scheme not in {"http", "https"}:
            raise ValueError("Only http and https URLs are supported")
        try:
            response = httpx.get(url, timeout=timeout, follow_redirects=True)
        except httpx.HTTPError as exc:  # pragma: no cover - network failure
            raise ValueError(f"Failed to download URL: {exc}") from exc
        if response.status_code >= 400:
            raise ValueError(f"Failed to download URL (status {response.status_code})")
        content = response.content
        if not content:
            raise ValueError("URL returned no content")

        content_type = response.headers.get("content-type")
        if content_type:
            content_type = content_type.split(";", 1)[0].strip()

        filename = Path(parsed.path).name or "document"
        if not Path(filename).suffix and content_type:
            extension = mimetypes.guess_extension(content_type.split(";", 1)[0])
            if extension:
                filename = f"{filename}{extension}"

        return self.ingest_bytes(
            project_id,
            filename=filename,
            data=content,
            content_type=content_type,
        )

    def _build_records(
        self,
        doc_id: str,
        filename: str,
        chunks: Sequence[Chunk],
        embeddings: Sequence[Sequence[float]],
        project_id: str,
        content_type: str | None,
        title: str | None,
    ) -> tuple[list[VectorRecord], list[dict[str, str]]]:
        records: list[VectorRecord] = []
        fts_entries: list[dict[str, str]] = []
        ingested_at = self._now()
        for index, (chunk, vector) in enumerate(zip(chunks, embeddings, strict=True)):
            chunk_id = f"{doc_id}:{index}"
            section_title = chunk.section_title or title
            payload = {
                "text": chunk.text,
                "project_id": project_id,
                "doc_id": doc_id,
                "chunk_index": index,
                "chunk_id": chunk_id,
                "source": filename,
                "content_type": content_type,
                "title": section_title,
                "heading_path": list(chunk.heading_path),
                "summary": chunk.summary,
                "language": chunk.language,
                "token_count": chunk.token_count,
                "char_count": chunk.char_count,
                "section_path": " / ".join(chunk.heading_path) if chunk.heading_path else section_title,
                "ingested_at": ingested_at,
            }
            records.append(
                VectorRecord(
                    id=uuid4().hex,
                    vector=list(vector),
                    payload=payload,
                )
            )
            fts_entries.append(
                {
                    "chunk_id": chunk_id,
                    "text": chunk.text,
                    "title": section_title or "",
                    "metadata": {
                        "source": filename,
                        "heading_path": list(chunk.heading_path),
                        "content_type": content_type,
                        "summary": chunk.summary,
                        "language": chunk.language,
                        "token_count": chunk.token_count,
                        "section_path": " / ".join(chunk.heading_path) if chunk.heading_path else section_title,
                        "ingested_at": ingested_at,
                    },
                }
            )
        return records, fts_entries

    @staticmethod
    def _extract_document(filename: str, data: bytes, content_type: str | None) -> ExtractedDocument:
        suffix = Path(filename).suffix.lower()
        guessed_type = content_type or mimetypes.guess_type(filename)[0]

        if suffix in {".md", ".markdown"}:
            text = data.decode("utf-8", errors="ignore")
            title = DocumentIngestor._derive_markdown_title(text)
            return ExtractedDocument(text=text, title=title, content_type=guessed_type or "text/markdown")

        if suffix in {".txt", ""} or (
            guessed_type
            and guessed_type.startswith("text/")
            and "html" not in guessed_type.lower()
        ):
            text = data.decode("utf-8", errors="ignore")
            title = DocumentIngestor._derive_plain_title(text)
            return ExtractedDocument(text=text, title=title, content_type=guessed_type or "text/plain")

        if suffix == ".pdf":
            if PdfReader is None:
                raise ValueError("PDF support requires the 'pypdf' package")
            try:
                reader = PdfReader(io.BytesIO(data))
                texts = [page.extract_text() or "" for page in reader.pages]
                metadata_title = None
                try:  # pragma: no cover - optional metadata
                    metadata_title = getattr(reader, "metadata", None)
                    if metadata_title and getattr(metadata_title, "title", None):
                        metadata_title = metadata_title.title
                    else:
                        metadata_title = None
                except Exception:
                    metadata_title = None
            except Exception as exc:  # pragma: no cover - parsing errors
                raise ValueError(f"Failed to extract text from PDF: {exc}") from exc
            combined = "\n\n".join(filter(None, texts))
            if not combined.strip():
                combined = data.decode("utf-8", errors="ignore")
            title = metadata_title or DocumentIngestor._derive_plain_title(combined)
            return ExtractedDocument(text=combined, title=title, content_type=guessed_type or "application/pdf")

        if suffix == ".docx":
            if DocxDocument is None:
                raise ValueError("DOCX support requires the 'python-docx' package")
            try:
                document = DocxDocument(io.BytesIO(data))
            except Exception as exc:  # pragma: no cover - parsing errors
                raise ValueError(f"Failed to extract text from DOCX: {exc}") from exc
            paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            text = "\n\n".join(paragraphs)
            core_title = None
            try:
                core_title = document.core_properties.title
            except Exception:  # pragma: no cover
                core_title = None
            base_title = core_title or (paragraphs[0] if paragraphs else None)
            if base_title:
                base_title = base_title.strip()
            return ExtractedDocument(
                text=text,
                title=base_title or DocumentIngestor._derive_plain_title(text),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        if suffix in {".html", ".htm"} or (guessed_type and "html" in guessed_type):
            soup_cls = _require_beautifulsoup()
            if soup_cls is None:
                raise ValueError("HTML support requires the 'beautifulsoup4' package")
            soup = soup_cls(data, "html.parser")
            for tag in soup(["script", "style"]):
                tag.extract()
            title = None
            if soup.title and soup.title.string:
                title = soup.title.string.strip()
            text = soup.get_text(separator="\n")
            return ExtractedDocument(
                text=text,
                title=title or DocumentIngestor._derive_plain_title(text),
                content_type="text/html",
            )

        # Fallback: treat as UTF-8 text
        text = data.decode("utf-8", errors="ignore")
        return ExtractedDocument(
            text=text,
            title=DocumentIngestor._derive_plain_title(text),
            content_type=guessed_type or "text/plain",
        )

    @staticmethod
    def _derive_markdown_title(text: str) -> str | None:
        for line in text.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("#"):
                return stripped.lstrip("#").strip()[:160] or None
            return stripped[:160]
        return None

    @staticmethod
    def _derive_plain_title(text: str) -> str | None:
        for line in text.splitlines():
            stripped = line.strip()
            if stripped:
                return stripped[:160]
        return None

    @staticmethod
    def _now() -> str:
        return datetime.now(timezone.utc).isoformat()


__all__ = [
    "DocumentIngestor",
    "ProjectVectorStoreManager",
    "IngestionResult",
    "IngestionJob",
    "IngestionJobManager",
]
